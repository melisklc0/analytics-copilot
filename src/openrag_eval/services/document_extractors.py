from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from markdown_it import MarkdownIt
from pypdf import PdfReader

from openrag_eval.core.exceptions import (
    DocumentTextExtractionError,
    UnsupportedDocumentFormatError,
)
from openrag_eval.schemas.document import DocumentFormat


class ParsedDocument:
    """Text and metadata inferred from an uploaded document file."""

    def __init__(self, title: str, content: str, format: DocumentFormat) -> None:
        self.title = title
        self.content = content
        self.format = format


class DocumentExtractorService:
    """Parse supported document uploads into text for ingestion."""

    _formats_by_extension = {
        ".txt": DocumentFormat.PLAIN_TEXT,
        ".md": DocumentFormat.MARKDOWN,
        ".markdown": DocumentFormat.MARKDOWN,
        ".pdf": DocumentFormat.PDF,
        ".docx": DocumentFormat.DOCX,
    }

    def __init__(self) -> None:
        self._markdown_reader = MarkdownIt("commonmark")

    async def parse_upload(self, *, filename: str, data: bytes) -> ParsedDocument:
        path = Path(filename)
        document_format = self._formats_by_extension.get(path.suffix.lower())
        if document_format is None:
            raise UnsupportedDocumentFormatError()

        content = self._extract_text(document_format, data).strip()
        if not content:
            raise DocumentTextExtractionError("Uploaded document has no extractable text")

        title = path.stem or "uploaded-document"
        return ParsedDocument(title=title, content=content, format=document_format)

    def _extract_text(self, document_format: DocumentFormat, data: bytes) -> str:
        if document_format is DocumentFormat.PLAIN_TEXT:
            return self._decode_text(data)
        if document_format is DocumentFormat.MARKDOWN:
            return self._extract_markdown_text(data)
        if document_format is DocumentFormat.PDF:
            return self._extract_pdf_text(data)
        if document_format is DocumentFormat.DOCX:
            return self._extract_docx_text(data)

        raise UnsupportedDocumentFormatError()

    def _decode_text(self, data: bytes) -> str:
        try:
            return data.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise DocumentTextExtractionError(
                "Uploaded text document must be valid UTF-8",
            ) from exc

    def _extract_markdown_text(self, data: bytes) -> str:
        content = self._decode_text(data)
        try:
            self._markdown_reader.parse(content)
        except Exception as exc:
            raise DocumentTextExtractionError(
                "Could not parse uploaded Markdown document",
            ) from exc
        return content

    def _extract_pdf_text(self, data: bytes) -> str:
        try:
            reader = PdfReader(BytesIO(data))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise DocumentTextExtractionError() from exc

    def _extract_docx_text(self, data: bytes) -> str:
        try:
            document = DocxDocument(BytesIO(data))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception as exc:
            raise DocumentTextExtractionError() from exc


document_extractor_service = DocumentExtractorService()
